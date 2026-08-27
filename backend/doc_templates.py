"""
Document template engine — generates PDF and DOCX documents from templates.

Templates available:
  - rental_agreement: Residential rental agreement (Indian context)
  - nda: Non-disclosure agreement
  - will: Last will and testament
  - employment_contract: Employment agreement
  - loan_agreement: Personal loan agreement
  - power_of_attorney: General power of attorney
  - sale_deed: Property sale deed (summary)
  - partnership_deed: Partnership agreement

Each template has:
  - fields: list of {key, label, type, required, default, placeholder}
  - generate_pdf(data) -> bytes
  - generate_docx(data) -> bytes
"""

import io
import os
import logging
from datetime import datetime
from typing import Optional

from reportlab.lib.pagesizes import A4
from reportlab.lib.units import inch, cm
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, PageBreak,
    Table, TableStyle, HRFlowable
)
from reportlab.lib import colors

logger = logging.getLogger(__name__)


# ── Template definitions ─────────────────────────────────────────────────────

TEMPLATES = {
    "rental_agreement": {
        "name": "Rental Agreement",
        "description": "Residential rental agreement between landlord and tenant",
        "fields": [
            {"key": "landlord_name", "label": "Landlord Name", "type": "text", "required": True, "placeholder": "Mr. Rajesh Kumar"},
            {"key": "landlord_address", "label": "Landlord Address", "type": "textarea", "required": True, "placeholder": "123 MG Road, Bangalore, KA 560001"},
            {"key": "tenant_name", "label": "Tenant Name", "type": "text", "required": True, "placeholder": "Ms. Priya Sharma"},
            {"key": "tenant_address", "label": "Tenant Current Address", "type": "textarea", "required": True, "placeholder": "456 Brigade Road, Bangalore, KA 560025"},
            {"key": "property_address", "label": "Property Address (rented)", "type": "textarea", "required": True, "placeholder": "Flat 302, Prestige Apartments, Indiranagar, Bangalore 560038"},
            {"key": "monthly_rent", "label": "Monthly Rent (Rs)", "type": "number", "required": True, "placeholder": "25000"},
            {"key": "security_deposit", "label": "Security Deposit (Rs)", "type": "number", "required": True, "placeholder": "50000"},
            {"key": "lease_start", "label": "Lease Start Date", "type": "date", "required": True, "placeholder": "2024-01-01"},
            {"key": "lease_end", "label": "Lease End Date", "type": "date", "required": True, "placeholder": "2024-12-31"},
            {"key": "notice_period", "label": "Notice Period (months)", "type": "number", "required": False, "default": "1", "placeholder": "1"},
            {"key": "city", "label": "City (for jurisdiction)", "type": "text", "required": True, "placeholder": "Bangalore"},
        ],
    },
    "nda": {
        "name": "Non-Disclosure Agreement",
        "description": "Confidentiality agreement between two parties",
        "fields": [
            {"key": "discloser_name", "label": "Disclosing Party Name", "type": "text", "required": True, "placeholder": "ABC Technologies Pvt Ltd"},
            {"key": "discloser_address", "label": "Disclosing Party Address", "type": "textarea", "required": True, "placeholder": "Tech Park, Whitefield, Bangalore 560066"},
            {"key": "receiver_name", "label": "Receiving Party Name", "type": "text", "required": True, "placeholder": "XYZ Solutions Inc"},
            {"key": "receiver_address", "label": "Receiving Party Address", "type": "textarea", "required": True, "placeholder": "100 Market St, San Francisco, CA 94103"},
            {"key": "purpose", "label": "Purpose of Disclosure", "type": "textarea", "required": True, "placeholder": "Evaluating a potential business partnership for software development"},
            {"key": "duration_months", "label": "Confidentiality Period (months)", "type": "number", "required": False, "default": "24", "placeholder": "24"},
            {"key": "jurisdiction", "label": "Jurisdiction City", "type": "text", "required": True, "placeholder": "Bangalore"},
            {"key": "agreement_date", "label": "Agreement Date", "type": "date", "required": True, "placeholder": "2024-01-15"},
        ],
    },
    "will": {
        "name": "Last Will and Testament",
        "description": "Legal will for distribution of assets",
        "fields": [
            {"key": "testator_name", "label": "Full Name of Testator", "type": "text", "required": True, "placeholder": "Mr. Rajesh Kumar"},
            {"key": "testator_address", "label": "Address", "type": "textarea", "required": True, "placeholder": "123 MG Road, Bangalore, KA 560001"},
            {"key": "testator_age", "label": "Age", "type": "number", "required": True, "placeholder": "45"},
            {"key": "testator_aadhaar", "label": "Aadhaar Number (optional)", "type": "text", "required": False, "placeholder": "XXXX-XXXX-1234"},
            {"key": "executor_name", "label": "Executor Name", "type": "text", "required": True, "placeholder": "Mrs. Sunita Kumar"},
            {"key": "executor_address", "label": "Executor Address", "type": "textarea", "required": True, "placeholder": "123 MG Road, Bangalore, KA 560001"},
            {"key": "beneficiaries", "label": "Beneficiaries & Distribution", "type": "textarea", "required": True, "placeholder": "1. Mrs. Sunita Kumar (wife) - 50% of all property\n2. Master Arjun Kumar (son) - 25% of all property\n3. Ms. Anjali Kumar (daughter) - 25% of all property"},
            {"key": "specific_bequests", "label": "Specific Bequests (optional)", "type": "textarea", "required": False, "placeholder": "My gold necklace to my daughter Anjali. My car to my son Arjun."},
            {"key": "witness1_name", "label": "Witness 1 Name", "type": "text", "required": True, "placeholder": "Mr. Amit Verma"},
            {"key": "witness1_address", "label": "Witness 1 Address", "type": "textarea", "required": True, "placeholder": "789 Residency Rd, Bangalore 560001"},
            {"key": "witness2_name", "label": "Witness 2 Name", "type": "text", "required": True, "placeholder": "Ms. Kavita Nair"},
            {"key": "witness2_address", "label": "Witness 2 Address", "type": "textarea", "required": True, "placeholder": "321 Park Ave, Bangalore 560034"},
            {"key": "city", "label": "City", "type": "text", "required": True, "placeholder": "Bangalore"},
            {"key": "will_date", "label": "Date of Will", "type": "date", "required": True, "placeholder": "2024-01-15"},
        ],
    },
    "employment_contract": {
        "name": "Employment Agreement",
        "description": "Employment contract between employer and employee",
        "fields": [
            {"key": "employer_name", "label": "Employer Name", "type": "text", "required": True, "placeholder": "ABC Technologies Pvt Ltd"},
            {"key": "employer_address", "label": "Employer Address", "type": "textarea", "required": True, "placeholder": "Tech Park, Whitefield, Bangalore 560066"},
            {"key": "employee_name", "label": "Employee Name", "type": "text", "required": True, "placeholder": "Mr. John Doe"},
            {"key": "employee_address", "label": "Employee Address", "type": "textarea", "required": True, "placeholder": "456 Brigade Road, Bangalore 560025"},
            {"key": "designation", "label": "Designation", "type": "text", "required": True, "placeholder": "Senior Software Engineer"},
            {"key": "salary", "label": "Annual CTC (Rs)", "type": "number", "required": True, "placeholder": "1500000"},
            {"key": "start_date", "label": "Start Date", "type": "date", "required": True, "placeholder": "2024-02-01"},
            {"key": "probation_months", "label": "Probation Period (months)", "type": "number", "required": False, "default": "6", "placeholder": "6"},
            {"key": "notice_months", "label": "Notice Period (months)", "type": "number", "required": False, "default": "2", "placeholder": "2"},
            {"key": "city", "label": "City", "type": "text", "required": True, "placeholder": "Bangalore"},
        ],
    },
    "loan_agreement": {
        "name": "Loan Agreement",
        "description": "Personal loan agreement between lender and borrower",
        "fields": [
            {"key": "lender_name", "label": "Lender Name", "type": "text", "required": True, "placeholder": "Mr. Suresh Patel"},
            {"key": "lender_address", "label": "Lender Address", "type": "textarea", "required": True, "placeholder": "123 Lake View, Mumbai 400001"},
            {"key": "borrower_name", "label": "Borrower Name", "type": "text", "required": True, "placeholder": "Mr. Ramesh Gupta"},
            {"key": "borrower_address", "label": "Borrower Address", "type": "textarea", "required": True, "placeholder": "456 Hill Road, Mumbai 400002"},
            {"key": "principal", "label": "Principal Amount (Rs)", "type": "number", "required": True, "placeholder": "500000"},
            {"key": "interest_rate", "label": "Interest Rate (% per annum)", "type": "number", "required": True, "placeholder": "12"},
            {"key": "loan_date", "label": "Loan Date", "type": "date", "required": True, "placeholder": "2024-01-15"},
            {"key": "repayment_date", "label": "Repayment Date", "type": "date", "required": True, "placeholder": "2025-01-15"},
            {"key": "repayment_mode", "label": "Repayment Mode", "type": "text", "required": False, "default": "Lump sum on due date", "placeholder": "Monthly EMI / Lump sum"},
            {"key": "city", "label": "City", "type": "text", "required": True, "placeholder": "Mumbai"},
        ],
    },
    "power_of_attorney": {
        "name": "Power of Attorney",
        "description": "General power of attorney authorizing an agent",
        "fields": [
            {"key": "principal_name", "label": "Principal Name", "type": "text", "required": True, "placeholder": "Mr. Rajesh Kumar"},
            {"key": "principal_address", "label": "Principal Address", "type": "textarea", "required": True, "placeholder": "123 MG Road, Bangalore 560001"},
            {"key": "agent_name", "label": "Agent Name", "type": "text", "required": True, "placeholder": "Mrs. Sunita Kumar"},
            {"key": "agent_address", "label": "Agent Address", "type": "textarea", "required": True, "placeholder": "123 MG Road, Bangalore 560001"},
            {"key": "scope", "label": "Scope of Authority", "type": "textarea", "required": False, "default": "General - to manage property, bank accounts, sign documents, and represent in all matters", "placeholder": "General or specific scope"},
            {"key": "valid_from", "label": "Valid From", "type": "date", "required": True, "placeholder": "2024-01-15"},
            {"key": "valid_until", "label": "Valid Until (optional)", "type": "date", "required": False, "placeholder": "2025-01-15"},
            {"key": "city", "label": "City", "type": "text", "required": True, "placeholder": "Bangalore"},
        ],
    },
    "partnership_deed": {
        "name": "Partnership Deed",
        "description": "Partnership agreement between two or more partners",
        "fields": [
            {"key": "firm_name", "label": "Firm Name", "type": "text", "required": True, "placeholder": "Kumar & Sharma Associates"},
            {"key": "partner1_name", "label": "Partner 1 Name", "type": "text", "required": True, "placeholder": "Mr. Arvind Kumar"},
            {"key": "partner2_name", "label": "Partner 2 Name", "type": "text", "required": True, "placeholder": "Mr. Deepak Sharma"},
            {"key": "partner1_address", "label": "Partner 1 Address", "type": "textarea", "required": True, "placeholder": "123 MG Road, Bangalore 560001"},
            {"key": "partner2_address", "label": "Partner 2 Address", "type": "textarea", "required": True, "placeholder": "456 Brigade Road, Bangalore 560025"},
            {"key": "business_address", "label": "Business Address", "type": "textarea", "required": True, "placeholder": "789 Commercial Street, Bangalore 560002"},
            {"key": "business_nature", "label": "Nature of Business", "type": "text", "required": True, "placeholder": "IT Consulting and Software Development"},
            {"key": "capital1", "label": "Partner 1 Capital (Rs)", "type": "number", "required": True, "placeholder": "500000"},
            {"key": "capital2", "label": "Partner 2 Capital (Rs)", "type": "number", "required": True, "placeholder": "500000"},
            {"key": "profit_ratio", "label": "Profit Sharing Ratio", "type": "text", "required": False, "default": "50:50", "placeholder": "50:50"},
            {"key": "commencement_date", "label": "Commencement Date", "type": "date", "required": True, "placeholder": "2024-01-01"},
            {"key": "city", "label": "City", "type": "text", "required": True, "placeholder": "Bangalore"},
        ],
    },
    "sale_deed": {
        "name": "Sale Deed (Summary)",
        "description": "Property sale deed summary between seller and buyer",
        "fields": [
            {"key": "seller_name", "label": "Seller Name", "type": "text", "required": True, "placeholder": "Mr. Mohan Reddy"},
            {"key": "seller_address", "label": "Seller Address", "type": "textarea", "required": True, "placeholder": "123 Jayanagar, Bangalore 560011"},
            {"key": "buyer_name", "label": "Buyer Name", "type": "text", "required": True, "placeholder": "Mr. Krishnan Nair"},
            {"key": "buyer_address", "label": "Buyer Address", "type": "textarea", "required": True, "placeholder": "456 Koramangala, Bangalore 560034"},
            {"key": "property_details", "label": "Property Details", "type": "textarea", "required": True, "placeholder": "Site No. 42, measuring 1200 sq ft, at Layout, Bangalore, bearing Khata No. 1234"},
            {"key": "sale_price", "label": "Sale Consideration (Rs)", "type": "number", "required": True, "placeholder": "7500000"},
            {"key": "sale_date", "label": "Date of Sale", "type": "date", "required": True, "placeholder": "2024-01-20"},
            {"key": "city", "label": "City", "type": "text", "required": True, "placeholder": "Bangalore"},
        ],
    },
}


# ── Utility functions ───────────────────────────────────────────────────────

def _format_date(date_str: str) -> str:
    """Format YYYY-MM-DD to 'DD day of Month, Year'."""
    if not date_str:
        return "____________"
    try:
        dt = datetime.strptime(date_str, "%Y-%m-%d")
        return dt.strftime("%d day of %B, %Y")
    except Exception:
        return date_str


def _format_amount(amount) -> str:
    """Format a number as Indian rupees."""
    try:
        n = float(amount)
        return f"Rs. {n:,.0f} ({num_to_words(n)} rupees)"
    except Exception:
        return str(amount)


def num_to_words(n):
    """Convert number to words (Indian system)."""
    try:
        n = int(n)
    except Exception:
        return ""
    if n == 0:
        return "zero"
    ones = ["", "one", "two", "three", "four", "five", "six", "seven", "eight", "nine",
            "ten", "eleven", "twelve", "thirteen", "fourteen", "fifteen", "sixteen",
            "seventeen", "eighteen", "nineteen"]
    tens = ["", "", "twenty", "thirty", "forty", "fifty", "sixty", "seventy", "eighty", "ninety"]

    def two_digit(num):
        if num < 20:
            return ones[num]
        return tens[num // 10] + (" " + ones[num % 10] if num % 10 else "")

    def three_digit(num):
        h = num // 100
        r = num % 100
        if h:
            return ones[h] + " hundred" + (" " + two_digit(r) if r else "")
        return two_digit(r)

    if n >= 10000000:
        return num_to_words(n // 10000000) + " crore " + num_to_words(n % 10000000)
    elif n >= 100000:
        return num_to_words(n // 100000) + " lakh " + num_to_words(n % 100000)
    elif n >= 1000:
        return num_to_words(n // 1000) + " thousand " + num_to_words(n % 1000)
    else:
        return three_digit(n).strip()


# ── PDF Generation ───────────────────────────────────────────────────────────

def _pdf_styles():
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="DocTitle", parent=styles["Title"], fontSize=18, spaceAfter=20, alignment=TA_CENTER, fontName="Helvetica-Bold"))
    styles.add(ParagraphStyle(name="DocSubtitle", parent=styles["Normal"], fontSize=10, spaceAfter=16, alignment=TA_CENTER, textColor=colors.grey))
    styles.add(ParagraphStyle(name="DocBody", parent=styles["Normal"], fontSize=11, leading=16, spaceAfter=10, alignment=TA_JUSTIFY))
    styles.add(ParagraphStyle(name="DocHeading", parent=styles["Heading2"], fontSize=12, spaceBefore=14, spaceAfter=8, fontName="Helvetica-Bold"))
    return styles


def generate_pdf(template_id: str, data: dict) -> bytes:
    template = TEMPLATES.get(template_id)
    if not template:
        raise ValueError(f"Unknown template: {template_id}")
    d = {f["key"]: data.get(f["key"], f.get("default", "")) for f in template["fields"]}
    styles = _pdf_styles()
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, leftMargin=1*inch, rightMargin=1*inch, topMargin=1*inch, bottomMargin=1*inch)
    story = []

    story.append(Paragraph(template["name"].upper(), styles["DocTitle"]))
    story.append(Paragraph(f"Generated on {datetime.now().strftime('%d %B %Y')}", styles["DocSubtitle"]))
    story.append(HRFlowable(width="100%", thickness=1, color=colors.grey))
    story.append(Spacer(1, 20))

    content_map = {
        "rental_agreement": _pdf_rental,
        "nda": _pdf_nda,
        "will": _pdf_will,
        "employment_contract": _pdf_employment,
        "loan_agreement": _pdf_loan,
        "power_of_attorney": _pdf_poa,
        "partnership_deed": _pdf_partnership,
        "sale_deed": _pdf_sale,
    }
    generator = content_map.get(template_id)
    if generator:
        story.extend(generator(d, styles))

    story.append(Spacer(1, 40))
    story.append(HRFlowable(width="100%", thickness=0.5, color=colors.grey))
    story.append(Spacer(1, 20))
    story.append(Paragraph("<b>Signature(s)</b>", styles["DocHeading"]))
    story.append(Spacer(1, 20))

    sig_parties = _get_sig_parties(template_id, d)
    if sig_parties:
        cols = len(sig_parties)
        sig_data = [["" for _ in range(cols)], [Paragraph(f"<b>{p}</b>", styles["Normal"]) for p in sig_parties]]
        col_w = 9360 / cols
        sig_table = Table(sig_data, colWidths=[col_w] * cols)
        sig_table.setStyle(TableStyle([("VALIGN", (0, 0), (-1, -1), "BOTTOM"), ("TOPPADDING", (0, 0), (-1, -1), 30)]))
        story.append(sig_table)

    story.append(Spacer(1, 30))
    story.append(Paragraph(
        "<i>This is a computer-generated document. Please review carefully before signing. "
        "Consult a legal professional to ensure compliance with applicable laws.</i>",
        ParagraphStyle("Disclaimer", parent=styles["Normal"], fontSize=8, textColor=colors.grey, alignment=TA_CENTER)
    ))

    doc.build(story)
    buf.seek(0)
    return buf.read()


def _get_sig_parties(template_id, d):
    parties_map = {
        "rental_agreement": [f"{d.get('landlord_name','')}\nLandlord", f"{d.get('tenant_name','')}\nTenant"],
        "nda": [f"{d.get('discloser_name','')}\nDisclosing Party", f"{d.get('receiver_name','')}\nReceiving Party"],
        "will": [f"{d.get('testator_name','')}\nTestator", f"Witness 1\n{d.get('witness1_name','')}", f"Witness 2\n{d.get('witness2_name','')}"],
        "employment_contract": [f"{d.get('employer_name','')}\nEmployer", f"{d.get('employee_name','')}\nEmployee"],
        "loan_agreement": [f"{d.get('lender_name','')}\nLender", f"{d.get('borrower_name','')}\nBorrower"],
        "power_of_attorney": [f"{d.get('principal_name','')}\nPrincipal", f"{d.get('agent_name','')}\nAgent"],
        "partnership_deed": [f"{d.get('partner1_name','')}\nPartner 1", f"{d.get('partner2_name','')}\nPartner 2"],
        "sale_deed": [f"{d.get('seller_name','')}\nSeller", f"{d.get('buyer_name','')}\nBuyer"],
    }
    return parties_map.get(template_id, ["Party 1", "Party 2"])


# ── PDF content generators ──────────────────────────────────────────────────

def _pdf_rental(d, s):
    return [
        Paragraph(f"THIS RENTAL AGREEMENT is made on <b>{_format_date(d.get('lease_start'))}</b> at <b>{d.get('city','')}</b>:", s["DocBody"]),
        Paragraph(f"BETWEEN <b>{d.get('landlord_name','')}</b>, residing at {d.get('landlord_address','')}, the <b>'LANDLORD'</b>;", s["DocBody"]),
        Paragraph(f"AND <b>{d.get('tenant_name','')}</b>, residing at {d.get('tenant_address','')}, the <b>'TENANT'</b>.", s["DocBody"]),
        Paragraph("WHEREAS the Landlord is the absolute owner of:", s["DocBody"]),
        Paragraph(f"<b>{d.get('property_address','')}</b>", s["DocBody"]),
        Paragraph("NOW THIS AGREEMENT WITNESSETH:", s["DocHeading"]),
        Paragraph(f"1. <b>Tenancy:</b> Monthly rent of <b>{_format_amount(d.get('monthly_rent'))}</b>.", s["DocBody"]),
        Paragraph(f"2. <b>Security Deposit:</b> <b>{_format_amount(d.get('security_deposit'))}</b> refundable.", s["DocBody"]),
        Paragraph(f"3. <b>Term:</b> From <b>{_format_date(d.get('lease_start'))}</b> to <b>{_format_date(d.get('lease_end'))}</b>.", s["DocBody"]),
        Paragraph("4. <b>Rent Payment:</b> By the 5th of each month.", s["DocBody"]),
        Paragraph(f"5. <b>Notice:</b> {d.get('notice_period','1')} month(s) written notice.", s["DocBody"]),
        Paragraph("6. <b>Use:</b> Residential purposes only. No sub-letting.", s["DocBody"]),
        Paragraph("7. <b>Maintenance:</b> Minor repairs by Tenant, major by Landlord.", s["DocBody"]),
        Paragraph("8. <b>Utilities:</b> Electricity, water borne by Tenant.", s["DocBody"]),
        Paragraph("9. <b>Inspection:</b> Landlord may inspect with prior intimation.", s["DocBody"]),
        Paragraph("10. <b>Termination:</b> Vacant possession in good condition on expiry.", s["DocBody"]),
        Paragraph("IN WITNESS WHEREOF, the parties have set their hands.", s["DocBody"]),
    ]

def _pdf_nda(d, s):
    return [
        Paragraph(f"THIS NON-DISCLOSURE AGREEMENT is entered into on <b>{_format_date(d.get('agreement_date'))}</b> at <b>{d.get('jurisdiction','')}</b>:", s["DocBody"]),
        Paragraph(f"BETWEEN <b>{d.get('discloser_name','')}</b>, at {d.get('discloser_address','')}, the <b>'DISCLOSING PARTY'</b>;", s["DocBody"]),
        Paragraph(f"AND <b>{d.get('receiver_name','')}</b>, at {d.get('receiver_address','')}, the <b>'RECEIVING PARTY'</b>.", s["DocBody"]),
        Paragraph("Purpose:", s["DocHeading"]),
        Paragraph(f"<b>{d.get('purpose','')}</b>", s["DocBody"]),
        Paragraph("NOW THEREFORE, the parties agree:", s["DocHeading"]),
        Paragraph("1. <b>Confidential Information:</b> Business plans, technical data, financial information, trade secrets.", s["DocBody"]),
        Paragraph("2. <b>Obligations:</b> Maintain confidentiality, no third-party disclosure, use solely for stated purpose.", s["DocBody"]),
        Paragraph("3. <b>Exclusions:</b> Already known, publicly available, or independently developed information.", s["DocBody"]),
        Paragraph(f"4. <b>Duration:</b> {d.get('duration_months','24')} months from disclosure.", s["DocBody"]),
        Paragraph("5. <b>Return:</b> Return or destroy all confidential information upon request.", s["DocBody"]),
        Paragraph("6. <b>No License:</b> No rights or licenses granted.", s["DocBody"]),
        Paragraph(f"7. <b>Governing Law:</b> Courts at {d.get('jurisdiction','')}.", s["DocBody"]),
        Paragraph("IN WITNESS WHEREOF, the parties have executed this agreement.", s["DocBody"]),
    ]

def _pdf_will(d, s):
    parts = [
        Paragraph(f"THIS IS THE LAST WILL AND TESTAMENT of me, <b>{d.get('testator_name','')}</b>, aged about <b>{d.get('testator_age','')}</b> years, residing at {d.get('testator_address','')}.", s["DocBody"]),
        Paragraph(f"Made at <b>{d.get('city','')}</b> on <b>{_format_date(d.get('will_date'))}</b>.", s["DocBody"]),
        Paragraph("I declare that I am making this will out of my own free will, in a sound state of mind. I revoke all previous wills.", s["DocBody"]),
        Paragraph(f"Executor: <b>{d.get('executor_name','')}</b>, {d.get('executor_address','')}.", s["DocBody"]),
        Paragraph("DISTRIBUTION OF ASSETS:", s["DocHeading"]),
    ]
    for line in (d.get('beneficiaries','')).split('\n'):
        if line.strip():
            parts.append(Paragraph(line.strip(), s["DocBody"]))
    if d.get('specific_bequests'):
        parts.append(Paragraph("SPECIFIC BEQUESTS:", s["DocHeading"]))
        parts.append(Paragraph(d.get('specific_bequests',''), s["DocBody"]))
    parts.append(Paragraph("I direct my Executor to pay all just debts, funeral expenses, and taxes before distribution.", s["DocBody"]))
    parts.append(Paragraph("SIGNED and DECLARED by the testator in the presence of:", s["DocBody"]))
    parts.append(Paragraph(f"Witness 1: {d.get('witness1_name','')} - {d.get('witness1_address','')}", s["DocBody"]))
    parts.append(Paragraph(f"Witness 2: {d.get('witness2_name','')} - {d.get('witness2_address','')}", s["DocBody"]))
    return parts

def _pdf_employment(d, s):
    return [
        Paragraph(f"THIS EMPLOYMENT AGREEMENT is executed on <b>{_format_date(d.get('start_date'))}</b> at <b>{d.get('city','')}</b>:", s["DocBody"]),
        Paragraph(f"Employer: <b>{d.get('employer_name','')}</b>, {d.get('employer_address','')}.", s["DocBody"]),
        Paragraph(f"Employee: <b>{d.get('employee_name','')}</b>, {d.get('employee_address','')}.", s["DocBody"]),
        Paragraph("Terms:", s["DocHeading"]),
        Paragraph(f"1. <b>Designation:</b> {d.get('designation','')}.", s["DocBody"]),
        Paragraph(f"2. <b>Compensation:</b> Annual CTC of <b>{_format_amount(d.get('salary'))}</b>.", s["DocBody"]),
        Paragraph(f"3. <b>Probation:</b> {d.get('probation_months','6')} months.", s["DocBody"]),
        Paragraph("4. <b>Duties:</b> As assigned by the Employer.", s["DocBody"]),
        Paragraph(f"5. <b>Notice Period:</b> {d.get('notice_months','2')} month(s).", s["DocBody"]),
        Paragraph("6. <b>Confidentiality:</b> Maintain confidentiality during and after employment.", s["DocBody"]),
        Paragraph("7. <b>IP:</b> All work product belongs to the Employer.", s["DocBody"]),
        Paragraph("IN WITNESS WHEREOF, the parties have set their hands.", s["DocBody"]),
    ]

def _pdf_loan(d, s):
    return [
        Paragraph(f"THIS LOAN AGREEMENT is made on <b>{_format_date(d.get('loan_date'))}</b> at <b>{d.get('city','')}</b>:", s["DocBody"]),
        Paragraph(f"Lender: <b>{d.get('lender_name','')}</b>, {d.get('lender_address','')}.", s["DocBody"]),
        Paragraph(f"Borrower: <b>{d.get('borrower_name','')}</b>, {d.get('borrower_address','')}.", s["DocBody"]),
        Paragraph("Terms:", s["DocHeading"]),
        Paragraph(f"1. <b>Loan Amount:</b> <b>{_format_amount(d.get('principal'))}</b>.", s["DocBody"]),
        Paragraph(f"2. <b>Interest:</b> {d.get('interest_rate','')}% per annum.", s["DocBody"]),
        Paragraph(f"3. <b>Repayment Date:</b> {_format_date(d.get('repayment_date'))}.", s["DocBody"]),
        Paragraph(f"4. <b>Repayment Mode:</b> {d.get('repayment_mode','Lump sum on due date')}.", s["DocBody"]),
        Paragraph("5. <b>Default:</b> Lender may take legal action on default.", s["DocBody"]),
        Paragraph("IN WITNESS WHEREOF, the parties have set their hands.", s["DocBody"]),
    ]

def _pdf_poa(d, s):
    parts = [
        Paragraph(f"THIS GENERAL POWER OF ATTORNEY is executed on <b>{_format_date(d.get('valid_from'))}</b> at <b>{d.get('city','')}</b>:", s["DocBody"]),
        Paragraph(f"Principal: <b>{d.get('principal_name','')}</b>, {d.get('principal_address','')}.", s["DocBody"]),
        Paragraph(f"Agent: <b>{d.get('agent_name','')}</b>, {d.get('agent_address','')}.", s["DocBody"]),
        Paragraph("Scope of Authority:", s["DocHeading"]),
        Paragraph(d.get('scope','General - to manage property, bank accounts, sign documents, and represent in all matters'), s["DocBody"]),
        Paragraph("1. Manage and administer all properties.", s["DocBody"]),
        Paragraph("2. Operate bank accounts and sign documents.", s["DocBody"]),
        Paragraph("3. Represent before courts, tribunals, and authorities.", s["DocBody"]),
        Paragraph("4. Execute and register documents on my behalf.", s["DocBody"]),
    ]
    if d.get('valid_until'):
        parts.append(Paragraph(f"5. Valid until <b>{_format_date(d.get('valid_until'))}</b>.", s["DocBody"]))
    parts.append(Paragraph("I agree to ratify all lawful acts done by the Agent.", s["DocBody"]))
    parts.append(Paragraph("IN WITNESS WHEREOF, I have set my hand.", s["DocBody"]))
    return parts

def _pdf_partnership(d, s):
    return [
        Paragraph(f"THIS DEED OF PARTNERSHIP is made on <b>{_format_date(d.get('commencement_date'))}</b> at <b>{d.get('city','')}</b>:", s["DocBody"]),
        Paragraph(f"Partner 1: <b>{d.get('partner1_name','')}</b>, {d.get('partner1_address','')}.", s["DocBody"]),
        Paragraph(f"Partner 2: <b>{d.get('partner2_name','')}</b>, {d.get('partner2_address','')}.", s["DocBody"]),
        Paragraph("Terms:", s["DocHeading"]),
        Paragraph(f"1. <b>Firm Name:</b> {d.get('firm_name','')}.", s["DocBody"]),
        Paragraph(f"2. <b>Business:</b> {d.get('business_nature','')} at {d.get('business_address','')}.", s["DocBody"]),
        Paragraph(f"3. <b>Capital:</b> Partner 1: {_format_amount(d.get('capital1'))}, Partner 2: {_format_amount(d.get('capital2'))}.", s["DocBody"]),
        Paragraph(f"4. <b>Profit Sharing:</b> {d.get('profit_ratio','50:50')}.", s["DocBody"]),
        Paragraph("5. <b>Management:</b> Equal rights for both partners.", s["DocBody"]),
        Paragraph("6. <b>Banking:</b> Operated by either or both jointly.", s["DocBody"]),
        Paragraph("7. <b>Dissolution:</b> By mutual consent.", s["DocBody"]),
        Paragraph("IN WITNESS WHEREOF, the parties have set their hands.", s["DocBody"]),
    ]

def _pdf_sale(d, s):
    return [
        Paragraph(f"THIS DEED OF SALE is executed on <b>{_format_date(d.get('sale_date'))}</b> at <b>{d.get('city','')}</b>:", s["DocBody"]),
        Paragraph(f"Seller: <b>{d.get('seller_name','')}</b>, {d.get('seller_address','')}.", s["DocBody"]),
        Paragraph(f"Buyer: <b>{d.get('buyer_name','')}</b>, {d.get('buyer_address','')}.", s["DocBody"]),
        Paragraph("Terms:", s["DocHeading"]),
        Paragraph(f"1. <b>Sale Consideration:</b> <b>{_format_amount(d.get('sale_price'))}</b>.", s["DocBody"]),
        Paragraph("2. The Seller conveys all rights, title, and interest to the Buyer.", s["DocBody"]),
        Paragraph("3. The property is free from all encumbrances.", s["DocBody"]),
        Paragraph("SCHEDULE OF PROPERTY:", s["DocHeading"]),
        Paragraph(f"<b>{d.get('property_details','')}</b>", s["DocBody"]),
        Paragraph("IN WITNESS WHEREOF, the parties have set their hands.", s["DocBody"]),
    ]


# ── DOCX Generation ─────────────────────────────────────────────────────────

def generate_docx(template_id: str, data: dict) -> bytes:
    try:
        from docx import Document as DocxDocument
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        return _generate_docx_fallback(template_id, data)

    template = TEMPLATES.get(template_id)
    if not template:
        raise ValueError(f"Unknown template: {template_id}")
    d = {f["key"]: data.get(f["key"], f.get("default", "")) for f in template["fields"]}
    doc = DocxDocument()

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    title = doc.add_heading(template["name"], level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Generated on {datetime.now().strftime('%d %B %Y')}")
    run.font.size = Pt(9)
    run.font.italic = True
    doc.add_paragraph()

    content_map = {
        "rental_agreement": _docx_rental,
        "nda": _docx_nda,
        "will": _docx_will,
        "employment_contract": _docx_employment,
        "loan_agreement": _docx_loan,
        "power_of_attorney": _docx_poa,
        "partnership_deed": _docx_partnership,
        "sale_deed": _docx_sale,
    }
    generator = content_map.get(template_id)
    if generator:
        generator(doc, d)

    doc.add_paragraph()
    doc.add_heading("Signature(s)", level=2)
    sig_parties = _get_sig_parties(template_id, d)
    for party in sig_parties:
        doc.add_paragraph("_" * 35)
        doc.add_paragraph(party.replace("\n", " - "))

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("This is a computer-generated document. Please review carefully before signing. "
                     "Consult a legal professional to ensure compliance with applicable laws.")
    run.font.size = Pt(8)
    run.font.italic = True

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def _docx_rental(doc, d):
    doc.add_paragraph(f"THIS RENTAL AGREEMENT is made on {_format_date(d.get('lease_start'))} at {d.get('city','')}:")
    doc.add_paragraph(f"Landlord: {d.get('landlord_name','')}, {d.get('landlord_address','')}")
    doc.add_paragraph(f"Tenant: {d.get('tenant_name','')}, {d.get('tenant_address','')}")
    doc.add_paragraph(f"Property: {d.get('property_address','')}")
    doc.add_heading("Terms:", level=2)
    doc.add_paragraph(f"1. Monthly Rent: {_format_amount(d.get('monthly_rent'))}")
    doc.add_paragraph(f"2. Security Deposit: {_format_amount(d.get('security_deposit'))}")
    doc.add_paragraph(f"3. Term: {_format_date(d.get('lease_start'))} to {_format_date(d.get('lease_end'))}")
    doc.add_paragraph("4. Rent by 5th of each month.")
    doc.add_paragraph(f"5. Notice: {d.get('notice_period','1')} month(s).")
    doc.add_paragraph("6. Residential use only. No sub-letting.")
    doc.add_paragraph("7. Minor repairs by Tenant, major by Landlord.")

def _docx_nda(doc, d):
    doc.add_paragraph(f"THIS NDA is entered into on {_format_date(d.get('agreement_date'))} at {d.get('jurisdiction','')}:")
    doc.add_paragraph(f"Disclosing Party: {d.get('discloser_name','')}, {d.get('discloser_address','')}")
    doc.add_paragraph(f"Receiving Party: {d.get('receiver_name','')}, {d.get('receiver_address','')}")
    doc.add_paragraph(f"Purpose: {d.get('purpose','')}")
    doc.add_heading("Terms:", level=2)
    doc.add_paragraph("1. Confidential Information: Business plans, technical data, trade secrets.")
    doc.add_paragraph("2. Obligations: Maintain confidentiality, no third-party disclosure.")
    doc.add_paragraph("3. Exclusions: Already known, public, or independently developed.")
    doc.add_paragraph(f"4. Duration: {d.get('duration_months','24')} months.")
    doc.add_paragraph("5. Return or destroy upon request.")
    doc.add_paragraph(f"6. Governing Law: Courts at {d.get('jurisdiction','')}.")

def _docx_will(doc, d):
    doc.add_paragraph(f"LAST WILL AND TESTAMENT of {d.get('testator_name','')}, aged {d.get('testator_age','')}, residing at {d.get('testator_address','')}.")
    doc.add_paragraph(f"Date: {_format_date(d.get('will_date'))} at {d.get('city','')}")
    doc.add_paragraph("I make this will freely, in sound mind. I revoke all previous wills.")
    doc.add_paragraph(f"Executor: {d.get('executor_name','')}, {d.get('executor_address','')}")
    doc.add_heading("Distribution of Assets:", level=2)
    for line in (d.get('beneficiaries','')).split('\n'):
        if line.strip():
            doc.add_paragraph(line.strip())
    if d.get('specific_bequests'):
        doc.add_heading("Specific Bequests:", level=2)
        doc.add_paragraph(d.get('specific_bequests',''))
    doc.add_paragraph("Executor to pay all debts and taxes before distribution.")
    doc.add_heading("Witnesses:", level=2)
    doc.add_paragraph(f"Witness 1: {d.get('witness1_name','')} - {d.get('witness1_address','')}")
    doc.add_paragraph(f"Witness 2: {d.get('witness2_name','')} - {d.get('witness2_address','')}")

def _docx_employment(doc, d):
    doc.add_paragraph(f"Employment Agreement dated {_format_date(d.get('start_date'))} at {d.get('city','')}")
    doc.add_paragraph(f"Employer: {d.get('employer_name','')}, {d.get('employer_address','')}")
    doc.add_paragraph(f"Employee: {d.get('employee_name','')}, {d.get('employee_address','')}")
    doc.add_heading("Terms:", level=2)
    doc.add_paragraph(f"1. Designation: {d.get('designation','')}")
    doc.add_paragraph(f"2. Annual CTC: {_format_amount(d.get('salary'))}")
    doc.add_paragraph(f"3. Probation: {d.get('probation_months','6')} months")
    doc.add_paragraph(f"4. Notice: {d.get('notice_months','2')} months")
    doc.add_paragraph("5. Confidentiality and IP clauses apply.")

def _docx_loan(doc, d):
    doc.add_paragraph(f"Loan Agreement dated {_format_date(d.get('loan_date'))} at {d.get('city','')}")
    doc.add_paragraph(f"Lender: {d.get('lender_name','')}, {d.get('lender_address','')}")
    doc.add_paragraph(f"Borrower: {d.get('borrower_name','')}, {d.get('borrower_address','')}")
    doc.add_heading("Terms:", level=2)
    doc.add_paragraph(f"1. Principal: {_format_amount(d.get('principal'))}")
    doc.add_paragraph(f"2. Interest: {d.get('interest_rate','')}% per annum")
    doc.add_paragraph(f"3. Repayment by: {_format_date(d.get('repayment_date'))}")
    doc.add_paragraph(f"4. Mode: {d.get('repayment_mode','Lump sum')}")

def _docx_poa(doc, d):
    doc.add_paragraph(f"Power of Attorney dated {_format_date(d.get('valid_from'))} at {d.get('city','')}")
    doc.add_paragraph(f"Principal: {d.get('principal_name','')}, {d.get('principal_address','')}")
    doc.add_paragraph(f"Agent: {d.get('agent_name','')}, {d.get('agent_address','')}")
    doc.add_heading("Scope:", level=2)
    doc.add_paragraph(d.get('scope','General'))
    doc.add_paragraph("Authorized to manage properties, operate bank accounts, represent before authorities.")
    if d.get('valid_until'):
        doc.add_paragraph(f"Valid until: {_format_date(d.get('valid_until'))}")

def _docx_partnership(doc, d):
    doc.add_paragraph(f"Partnership Deed dated {_format_date(d.get('commencement_date'))} at {d.get('city','')}")
    doc.add_paragraph(f"Partner 1: {d.get('partner1_name','')}, {d.get('partner1_address','')}")
    doc.add_paragraph(f"Partner 2: {d.get('partner2_name','')}, {d.get('partner2_address','')}")
    doc.add_heading("Terms:", level=2)
    doc.add_paragraph(f"1. Firm: {d.get('firm_name','')}")
    doc.add_paragraph(f"2. Business: {d.get('business_nature','')} at {d.get('business_address','')}")
    doc.add_paragraph(f"3. Capital: P1={_format_amount(d.get('capital1'))}, P2={_format_amount(d.get('capital2'))}")
    doc.add_paragraph(f"4. Profit: {d.get('profit_ratio','50:50')}")

def _docx_sale(doc, d):
    doc.add_paragraph(f"Sale Deed dated {_format_date(d.get('sale_date'))} at {d.get('city','')}")
    doc.add_paragraph(f"Seller: {d.get('seller_name','')}, {d.get('seller_address','')}")
    doc.add_paragraph(f"Buyer: {d.get('buyer_name','')}, {d.get('buyer_address','')}")
    doc.add_heading("Terms:", level=2)
    doc.add_paragraph(f"1. Sale Consideration: {_format_amount(d.get('sale_price'))}")
    doc.add_paragraph("2. Seller conveys all rights to Buyer.")
    doc.add_heading("Schedule of Property:", level=2)
    doc.add_paragraph(d.get('property_details',''))


def _generate_docx_fallback(template_id: str, data: dict) -> bytes:
    """Fallback: HTML-based document that opens in Word."""
    template = TEMPLATES.get(template_id, {})
    d = {f["key"]: data.get(f["key"], f.get("default", "")) for f in template.get("fields", [])}
    name = template.get("name", "Document")
    html = f"<html><head><meta charset='utf-8'></head><body><h1 style='text-align:center'>{name}</h1><hr/>"
    for f in template.get("fields", []):
        val = d.get(f["key"], "")
        if val:
            html += f"<p><b>{f['label']}:</b> {val}</p>"
    html += "<hr/><p><b>Signature(s)</b></p><p style='font-size:8pt;color:gray'>Computer-generated document.</p></body></html>"
    return html.encode("utf-8")


def get_template_list():
    """Return list of available templates for the API."""
    return [
        {"id": tid, "name": t["name"], "description": t["description"], "field_count": len(t["fields"]), "fields": t["fields"]}
        for tid, t in TEMPLATES.items()
    ]